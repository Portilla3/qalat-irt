"""
╔══════════════════════════════════════════════════════════════════════════════╗
║   SCRIPT_IRT_Universal_Word_Caracterizacion.py  —  v1.0                   ║
║   Genera informe Word de caracterización al ingreso (IRT1)                 ║
║   Compatible con cualquier país IRT                                        ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  CÓMO USAR:                                                                 ║
║  1. Abre un chat nuevo con Claude                                           ║
║  2. Sube DOS archivos:                                                      ║
║       • Este script                                                         ║
║       • La base Wide IRT (generada por SCRIPT_IRT_Universal_Wide)         ║
║  3. Escribe: "Ejecuta el script Word Caracterización IRT con esta base"    ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

import glob, os, unicodedata, io, warnings
import pandas as pd
import numpy as np
import matplotlib; matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
warnings.filterwarnings('ignore')

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6)
ACCENT=RGBColor(0x00,0xB0,0xF0); GRAY=RGBColor(0x59,0x59,0x59)
WHITE=RGBColor(0xFF,0xFF,0xFF)
C_NAVY='1F3864'; C_MID='2E75B6'; C_LIGHT='EEF4FB'
MC_MID='#2E75B6'; MC_LIGHT='#BDD7EE'; MC_ACCENT='#00B0F0'
PIE_COLS=['#2E75B6','#1F3864','#4472C4','#9DC3E6','#00B0F0','#538135','#BFBFBF','#C00000','#ED7D31']

_PAISES={
    'republica_dominicana':'República Dominicana','dominicana':'República Dominicana',
    'honduras':'Honduras','panama':'Panamá','panam':'Panamá',
    'el_salvador':'El Salvador','salvador':'El Salvador',
    'mexico':'México','mexic':'México','ecuador':'Ecuador',
    'peru':'Perú','argentina':'Argentina','colombia':'Colombia',
    'chile':'Chile','bolivia':'Bolivia','paraguay':'Paraguay',
    'uruguay':'Uruguay','venezuela':'Venezuela','guatemala':'Guatemala',
    'costa_rica':'Costa Rica','costarica':'Costa Rica','nicaragua':'Nicaragua',
}

def _norm(s): return unicodedata.normalize('NFD',str(s).lower()).encode('ascii','ignore').decode()
def _extraer_pais(fn):
    f=_norm(str(fn).replace('.','_'))
    for k,v in _PAISES.items():
        if k in f: return v
    return None

def _detectar_pais(wide_file):
    try:
        rs=pd.read_excel(wide_file,sheet_name='Resumen',header=None)
        for _,row in rs.iterrows():
            for v in row.tolist():
                p=_extraer_pais(str(v))
                if p: return p
    except: pass
    return _extraer_pais(os.path.basename(wide_file))

def auto_archivo_wide():
    candidatos=(
        glob.glob('/mnt/user-data/uploads/*IRT*Wide*.xlsx')+
        glob.glob('/mnt/user-data/uploads/*Wide*IRT*.xlsx')+
        glob.glob('/mnt/user-data/uploads/IRT_Base*.xlsx')+
        glob.glob('/home/claude/IRT_Base_Wide.xlsx'))
    if not candidatos: raise FileNotFoundError('No se encontró la base Wide IRT.')
    print(f'  → Base Wide: {os.path.basename(candidatos[0])}')
    return candidatos[0]

INPUT_FILE = None   # runner inyecta la ruta real
OUTPUT_FILE = None   # runner inyecta la ruta real
FILTRO_CENTRO = None   # runner inyecta el filtro si aplica

# Cargar datos
df=pd.read_excel(INPUT_FILE,sheet_name='Base Wide',header=1)
df.columns=[str(c) for c in df.columns]
cols=df.columns.tolist()

_col_centro=next((c for c in cols if any(x in _norm(c) for x in
                  ['codigo del centro','servicio de tratamiento',
                   'centro/ servicio','codigo centro'])),None)
if FILTRO_CENTRO and _col_centro:
    n_antes=len(df)
    df=df[df[_col_centro].astype(str).str.strip()==FILTRO_CENTRO].copy().reset_index(drop=True)
    print(f'\n⚑  Filtro: "{FILTRO_CENTRO}"  ({n_antes}→{len(df)} pacientes)')
    OUTPUT_FILE=f'/home/claude/IRT_Informe_Caracterizacion_{FILTRO_CENTRO}.docx'

mask1=df['Tiene_IRT1']=='Sí' if 'Tiene_IRT1' in cols else pd.Series([True]*len(df))
df1=df[mask1].copy().reset_index(drop=True); N=len(df1)
print(f'\n→ {N} pacientes IRT1 cargados')

PAIS=_detectar_pais(INPUT_FILE)
SERVICIO=(f'{PAIS}  —  Centro {FILTRO_CENTRO}' if FILTRO_CENTRO and PAIS
          else PAIS if PAIS else 'Servicio de Tratamiento')

MESES={1:'Enero',2:'Febrero',3:'Marzo',4:'Abril',5:'Mayo',6:'Junio',
       7:'Julio',8:'Agosto',9:'Septiembre',10:'Octubre',11:'Noviembre',12:'Diciembre'}
fecha_col=next((c for c in cols if 'fecha de administracion' in _norm(c)),None)
PERIODO='Período no determinado'
if fecha_col:
    fch=pd.to_datetime(df1[fecha_col],errors='coerce').dropna()
    anio=pd.Timestamp.now().year; fch=fch[(fch.dt.year>=anio-10)&(fch.dt.year<=anio+1)]
    if len(fch):
        f0,f1_=fch.min(),fch.max()
        PERIODO=(f'{MESES[f0.month]} {f0.year}' if f0.year==f1_.year and f0.month==f1_.month
                 else f'{MESES[f0.month]}–{MESES[f1_.month]} {f0.year}' if f0.year==f1_.year
                 else f'{MESES[f0.month]} {f0.year} – {MESES[f1_.month]} {f1_.year}')
print(f'  Servicio: {SERVICIO} | Período: {PERIODO}')

# Detección de columnas IRT1
def col1(kws):
    for c in cols:
        if not c.endswith('_IRT1'): continue
        if all(_norm(k) in _norm(c) for k in kws): return c
    return None

COL_SEXO=next((c for c in cols if _norm(c) in ['sexo','género','genero']),None)
COL_FN=next((c for c in cols if 'fecha de nacimiento' in _norm(c)),None)
COL_SP=col1(['sustancia','principal'])

SUST_NOMBRES={
    'Alcohol':['alcohol'],'Marihuana':['marihuana','cannabis'],
    'Heroína':['heroina'],'Cocaína':['cocain'],
    'Fentanilo':['fentanil'],'Inhalables':['inhalab'],
    'Metanfetamina':['metanfet','cristal'],'Crack':['crack'],
    'Pasta Base':['pasta base','pasta'],'Sedantes':['sedant','benzod'],
    'Opiáceos':['opiod','opiac'],'Tabaco':['tabaco','nicot'],
}
SUST_TOTAL1={}
for sust,kws in SUST_NOMBRES.items():
    for c in cols:
        if not c.endswith('_IRT1'): continue
        nc=_norm(c)
        if any(_norm(k) in nc for k in kws) and ('total' in nc or '(0-28)' in nc):
            SUST_TOTAL1[sust]=c; break
SUST_ACTIVAS=list(SUST_TOTAL1.keys())

COL_SPSI=col1(['salud','psicol']); COL_SFIS=col1(['salud','fis'])
COL_URG=next((c for c in cols if c.endswith('_IRT1') and '5)' in c and
              any(k in c.lower() for k in ['urgencia','hospitali','emergencia'])),None)
COL_ACC=next((c for c in cols if c.endswith('_IRT1') and '6)' in c and 'accidente' in c.lower()),None)

TRANS_DEF={'Robo / Hurto':'robo','Venta de sustancias':'venta',
           'Violencia a otras personas':'violencia',
           'Violencia intrafamiliar':'intraf','Detenido / Arrestado':'detenido'}
TRANS_COLS={n:next((c for c in cols if c.endswith('_IRT1') and kw in c.lower()),None)
            for n,kw in TRANS_DEF.items()}

REL_MAP={'Padre':['padre'],'Madre':['madre'],'Hijos':['hijos','hijo'],
         'Hermanos':['hermanos'],'Pareja':['pareja'],'Amigos':['amigos']}
REL_COLS={}
for vin,kws in REL_MAP.items():
    for c in cols:
        if not c.endswith('_IRT1') or ('14)' not in c and 'relaci' not in _norm(c)): continue
        if any(_norm(k) in _norm(c) for k in kws): REL_COLS[vin]=c; break

SAT_MAP={
    'Vida en general':[['16)'],['satisfac','vida']],
    'Lugar donde vive':[['17)'],['satisfac','lugar']],
    'Situación laboral':[['18)'],['satisfac','labor','educac']],
    'Tiempo libre':[['19)'],['satisfac','tiempo']],
    'Cap. económica':[['20)'],['satisfac','econom']],
}
SAT_COLS={}
for dim,(nums,kws) in SAT_MAP.items():
    for c in cols:
        if not c.endswith('_IRT1'): continue
        nc=_norm(c)
        if any(n in c for n in nums) and any(_norm(k) in nc for k in kws):
            SAT_COLS[dim]=c; break

# ── Cálculo de datos ──────────────────────────────────────────────────────────
print('\n→ Calculando datos...')
hoy=pd.Timestamp.now()

def norm_sust(s):
    if pd.isna(s) or str(s).strip() in ['0','']: return None
    s=_norm(str(s))
    if any(x in s for x in ['alcohol','cerveza','licor']): return 'Alcohol'
    if any(x in s for x in ['marihu','cannabis','marij']): return 'Marihuana'
    if any(x in s for x in ['crack','piedra','paco']): return 'Crack'
    if any(x in s for x in ['pasta base','pasta']): return 'Pasta Base'
    if any(x in s for x in ['cocain','perico']): return 'Cocaína'
    if any(x in s for x in ['fentanil']): return 'Fentanilo'
    if any(x in s for x in ['inhalab']): return 'Inhalables'
    if any(x in s for x in ['metanfet','cristal']): return 'Metanfetamina'
    if any(x in s for x in ['sedant','benzod']): return 'Sedantes'
    if any(x in s for x in ['heroina','opiod','morfin']): return 'Heroína'
    if any(x in s for x in ['tabaco','nicot']): return 'Tabaco'
    return 'Otra sustancia'

def prom1(col):
    if col is None: return np.nan,0
    v=pd.to_numeric(df1[col],errors='coerce').dropna()
    return (round(float(v.mean()),1) if len(v) else np.nan),len(v)

def sino1(col):
    if col is None: return None,None,None
    v=df1[col].dropna().astype(str).str.strip().str.lower()
    nv=len(v); nsi=int(v.isin(['sí','si']).sum())
    return nsi,round(nsi/nv*100,1) if nv else 0.0,nv

# Sexo
R_sexo={k:int(v) for k,v in df1[COL_SEXO].value_counts(dropna=True).items()} if COL_SEXO else {}
n_hombre=R_sexo.get('Hombre',R_sexo.get('hombre',R_sexo.get('H',0)))
n_mujer=R_sexo.get('Mujer',R_sexo.get('mujer',R_sexo.get('M',0)))
if n_hombre==0 and R_sexo:
    vals_sx=sorted(R_sexo.values(),reverse=True)
    n_hombre=vals_sx[0] if vals_sx else 0
    n_mujer=vals_sx[1] if len(vals_sx)>1 else 0
N_sx=n_hombre+n_mujer
pct_h=round(n_hombre/N_sx*100,1) if N_sx else 0
pct_m=round(n_mujer/N_sx*100,1) if N_sx else 0

# Edad
R_edad={}
if COL_FN:
    edades=((hoy-pd.to_datetime(df1[COL_FN],errors='coerce')).dt.days/365.25).dropna()
    edades=edades[(edades>=10)&(edades<=100)]
    if len(edades):
        bins=[0,17,25,35,45,55,200]; labs=['<18','18–25','26–35','36–45','46–55','56+']
        gr=pd.cut(edades,bins=bins,labels=labs).value_counts().reindex(labs,fill_value=0)
        R_edad={'grupos':{k:int(v) for k,v in gr.items()},
                'mean':round(edades.mean(),1),'std':round(edades.std(),1),
                'min':int(edades.min()),'max':int(edades.max())}

# Sustancia principal
sust_pp=df1[COL_SP].apply(norm_sust).dropna() if COL_SP else pd.Series(dtype=str)
R_sp=sust_pp.value_counts().to_dict() if len(sust_pp) else {}; N_sp=len(sust_pp)

# Días consumo
dias_pp=[]
for sust in SUST_ACTIVAS:
    col=SUST_TOTAL1.get(sust)
    if not col: continue
    msk=sust_pp==sust
    vals=pd.to_numeric(df1.loc[msk.reindex(df1.index,fill_value=False),col],errors='coerce').dropna()
    vals=vals[vals>0]
    if len(vals): dias_pp.append({'sust':sust,'prom':round(float(vals.mean()),1),'n':len(vals)})
dias_pp.sort(key=lambda x:-x['prom'])

cons_pct=[]
for sust,col in SUST_TOTAL1.items():
    v=pd.to_numeric(df1[col],errors='coerce').fillna(0); n_c=int((v>0).sum())
    if n_c: cons_pct.append({'sust':sust,'pct':round(n_c/N*100,1),'n':n_c})
cons_pct.sort(key=lambda x:-x['pct'])

dias_sust=[]
for sust,col in SUST_TOTAL1.items():
    v=pd.to_numeric(df1[col],errors='coerce'); sub=v[v>0].dropna()
    if len(sub): dias_sust.append({'sust':sust,'prom':round(float(sub.mean()),1),'n':len(sub)})
dias_sust.sort(key=lambda x:-x['prom'])

n_urg,pct_urg,nv_urg=sino1(COL_URG); n_acc,pct_acc,nv_acc=sino1(COL_ACC)

R_salud=[]
for nombre,col in [('Salud Psicológica (0–10)',COL_SPSI),('Salud Física (0–10)',COL_SFIS)]:
    m,nv=prom1(col)
    if not np.isnan(m): R_salud.append({'label':nombre,'prom':m,'nv':nv})

# Transgresión
n_tr=0; pct_tr=0; trans_tipos=[]
any_tr_series=[]
for nombre,col in TRANS_COLS.items():
    if col is None: continue
    v=df1[col].dropna().astype(str).str.lower()
    nsi=int(v.isin(['sí','si']).sum())
    if nsi==0:
        v2=pd.to_numeric(df1[col],errors='coerce').fillna(0); nsi=int((v2>0).sum())
    if nsi: trans_tipos.append({'label':nombre,'n':nsi,'pct':round(nsi/N*100,1)})
    any_tr_series.append(df1[col].astype(str).str.lower().isin(['sí','si']) if col else pd.Series([False]*N))
if any_tr_series:
    n_tr=int(pd.concat(any_tr_series,axis=1).any(axis=1).sum())
    pct_tr=round(n_tr/N*100,1) if N>0 else 0

# Relaciones
CATS_REL=['Excelente','Buena','Ni buena ni mala','Mala','Muy mala']
R_rel=[]
for vin,col in REL_COLS.items():
    vals=df1[col].dropna(); vals=vals[vals.astype(str).str.lower()!='no aplica']
    nv=len(vals)
    if nv==0: continue
    d={'vinculo':vin,'N':nv}
    for cat in CATS_REL:
        nc=int((vals.astype(str).str.lower()==cat.lower()).sum())
        d[cat]=round(nc/nv*100,1)
    R_rel.append(d)

# Satisfacción
R_sat=[]
for dim,col in SAT_COLS.items():
    m,nv=prom1(col)
    if not np.isnan(m): R_sat.append({'label':dim,'prom':m,'nv':nv})

print(f'  ✓ Datos calculados (N={N})')

# ── Helpers Word ──────────────────────────────────────────────────────────────
def set_cell_bg(cell,hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color); tcPr.append(shd)

def _keep_with_next(p):
    pPr=p._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepNext'))
    pPr.append(OxmlElement('w:keepLines'))

def _page_break(doc):
    from docx.enum.text import WD_BREAK
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)

def add_section_header(doc,num,title):
    tbl=doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
    c=tbl.rows[0].cells[0]; set_cell_bg(c,C_MID)
    p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    txt=f'  {num}. {title.upper()}' if num else f'  {title.upper()}'
    run=p.add_run(txt); run.font.name='Calibri'; run.font.size=Pt(11)
    run.font.bold=True; run.font.color.rgb=WHITE
    sp=doc.add_paragraph(); sp.paragraph_format.space_after=Pt(0); _keep_with_next(sp)

def add_subsection(doc,title):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2)
    run=p.add_run(title); run.bold=True; run.font.name='Calibri'; run.font.size=Pt(10.5)
    run.font.color.rgb=NAVY; _keep_with_next(p)

def add_picture_kwnext(doc,buf,width):
    doc.add_picture(buf,width=width)
    p=doc.paragraphs[-1]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; _keep_with_next(p)

def add_body(doc,text,italic=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
    run=p.add_run(text); run.font.name='Calibri'; run.font.size=Pt(10)
    run.font.italic=italic; run.font.color.rgb=RGBColor(0x33,0x33,0x33); return p

def add_note(doc,text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    run=p.add_run(text); run.font.name='Calibri'; run.font.size=Pt(8.5)
    run.font.italic=True; run.font.color.rgb=GRAY

def add_kpi_row(doc,kpis):
    tbl=doc.add_table(rows=2,cols=len(kpis)); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,(val,lbl) in enumerate(kpis):
        cv=tbl.rows[0].cells[j]; set_cell_bg(cv,C_LIGHT)
        p=cv.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(str(val)); run.font.name='Calibri'; run.font.size=Pt(22)
        run.font.bold=True; run.font.color.rgb=BLUE
        cl=tbl.rows[1].cells[j]; set_cell_bg(cl,C_LIGHT)
        p2=cl.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run2=p2.add_run(lbl); run2.font.name='Calibri'; run2.font.size=Pt(8.5)
        run2.font.color.rgb=GRAY
    doc.add_paragraph()

def fig_to_img(fig,width_cm=13):
    buf=io.BytesIO()
    fig.savefig(buf,format='png',dpi=150,bbox_inches='tight',facecolor='white')
    buf.seek(0); plt.close(fig); return buf,Cm(width_cm)

def _ax_style(ax,horiz=False):
    ax.set_facecolor('#F8FBFE')
    ax.grid(True,axis='x' if horiz else 'y',color='#E0E8F0',linewidth=0.7,zorder=0)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#CCCCCC'); ax.spines['bottom'].set_color('#CCCCCC')
    ax.tick_params(colors='#595959',labelsize=8)

# ── Gráficos ──────────────────────────────────────────────────────────────────
def g_sexo():
    if not R_sexo: return None
    labels=list(R_sexo.keys()); vals=list(R_sexo.values())
    fig,ax=plt.subplots(figsize=(5,4))
    w,t,at=ax.pie(vals,labels=None,colors=PIE_COLS[:len(vals)],
        autopct='%1.1f%%',startangle=90,pctdistance=0.72,
        wedgeprops={'edgecolor':'white','linewidth':1.5})
    for a in at: a.set_fontsize(10); a.set_fontweight('bold'); a.set_color('white')
    ax.legend(w,[f'{l} (n={v})' for l,v in zip(labels,vals)],
              loc='lower center',fontsize=9,frameon=False,bbox_to_anchor=(0.5,-0.12))
    ax.set_title(f'N = {sum(vals)}',fontsize=8.5,color='#595959',pad=4)
    fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_edad():
    if not R_edad: return None
    gr=R_edad['grupos']; labs=list(gr.keys()); vals=list(gr.values()); total=sum(vals)
    pcts=[round(v/total*100,1) if total else 0 for v in vals]
    fig,ax=plt.subplots(figsize=(6,3.5))
    bars=ax.bar(labs,pcts,color=MC_MID,width=0.6,zorder=3)
    for bar,v,p in zip(bars,vals,pcts):
        if v>0:
            ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.5,
                    f'{p}%\n(n={v})',ha='center',va='bottom',fontsize=7.5,fontweight='bold',color='#333')
    ax.set_ylim(0,max(pcts)*1.45+2 if pcts else 1)
    ax.set_ylabel('% del total',fontsize=8,color='#595959')
    _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_sp():
    if not R_sp: return None
    labs=list(R_sp.keys()); vals=list(R_sp.values())
    fig,ax=plt.subplots(figsize=(6,5.5))
    w,t,at=ax.pie(vals,labels=None,colors=PIE_COLS[:len(labs)],
        autopct='%1.1f%%',startangle=90,pctdistance=0.70,
        wedgeprops={'edgecolor':'white','linewidth':1.8},radius=0.92)
    for a in at: a.set_fontsize(9.5); a.set_fontweight('bold'); a.set_color('white')
    ax.legend(w,[f'{l} (n={v})' for l,v in zip(labs,vals)],
              loc='lower center',fontsize=8.0,frameon=False,bbox_to_anchor=(0.5,-0.28),ncol=1)
    ax.set_title(f'N válido = {sum(vals)}',fontsize=8.5,color='#595959',pad=6)
    ax.set_aspect('equal'); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_barras(datos,ylabel='',rotacion=0):
    if not datos: return None
    labs=[d.get('sust',d.get('label','')) for d in datos]
    vals=[d.get('prom',d.get('pct',0)) for d in datos]
    ns=[d.get('n',0) for d in datos]
    fig,ax=plt.subplots(figsize=(max(5,len(labs)*1.0),4))
    bars=ax.bar(labs,vals,color=MC_MID,width=0.55,zorder=3)
    for bar,v,n in zip(bars,vals,ns):
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.3,
                f'{v}\n(n={n})',ha='center',va='bottom',fontsize=8,fontweight='bold',color='#333')
    ax.set_ylim(0,max(vals)*1.4+2 if vals else 1)
    ax.set_ylabel(ylabel,fontsize=8,color='#595959')
    if rotacion: ax.set_xticklabels(labs,rotation=rotacion,ha='right')
    _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_salud_irt():
    if not R_salud: return None
    labs=[d['label'] for d in R_salud]; vals=[d['prom'] for d in R_salud]
    fig,ax=plt.subplots(figsize=(6,3.2))
    bars=ax.barh(labs,vals,color=MC_MID,height=0.45,zorder=3)
    for bar,v in zip(bars,vals):
        ax.text(bar.get_width()+0.1,bar.get_y()+bar.get_height()/2,
                str(v),va='center',fontsize=10,fontweight='bold',color='#333')
    ax.set_xlim(0,11.5); ax.axvline(x=5,color='#BFBFBF',linestyle='--',linewidth=0.8)
    ax.set_xlabel('Promedio (0–10)',fontsize=8,color='#595959')
    _ax_style(ax,horiz=True); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_transgresion_irt():
    fig,ax=plt.subplots(figsize=(5,3.5))
    bars=ax.bar(['Con transgresión','Sin transgresión'],
                [n_tr,N-n_tr],color=[MC_MID,MC_LIGHT],width=0.5,zorder=3)
    for bar,v,lab in zip(bars,[n_tr,N-n_tr],[f'{pct_tr}%',f'{round((N-n_tr)/N*100,1) if N>0 else 0}%']):
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.3,
                f'{v}\n({lab})',ha='center',va='bottom',fontsize=10,fontweight='bold',color='#333')
    ax.set_ylim(0,max(n_tr,N-n_tr)*1.4+2 if N>0 else 1)
    ax.set_ylabel(f'N personas (N={N})',fontsize=8,color='#595959')
    _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_relaciones():
    if not R_rel: return None
    vinculos=[d['vinculo'] for d in R_rel]
    pos=[d.get('Excelente',0)+d.get('Buena',0) for d in R_rel]
    neg=[d.get('Mala',0)+d.get('Muy mala',0) for d in R_rel]
    neu=[100-p-n for p,n in zip(pos,neg)]
    x=np.arange(len(vinculos))
    fig,ax=plt.subplots(figsize=(max(6,len(vinculos)*1.1),4))
    ax.bar(x,pos,color=MC_MID,width=0.55,zorder=3,label='Positiva')
    ax.bar(x,neu,bottom=pos,color='#D9D9D9',width=0.55,zorder=3,label='Neutral')
    ax.bar(x,neg,bottom=[p+n for p,n in zip(pos,neu)],color='#C00000',width=0.55,zorder=3,label='Negativa',alpha=0.8)
    for i,p in enumerate(pos):
        if p>5: ax.text(i,p/2,f'{p:.0f}%',ha='center',va='center',fontsize=8,fontweight='bold',color='white')
    ax.set_xticks(x); ax.set_xticklabels(vinculos,fontsize=9)
    ax.set_ylim(0,115); ax.set_ylabel('% personas',fontsize=8,color='#595959')
    ax.legend(fontsize=8,frameon=False,loc='upper right',ncol=3)
    _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_satisfaccion():
    if not R_sat: return None
    labs=[d['label'] for d in R_sat]; vals=[d['prom'] for d in R_sat]
    fig,ax=plt.subplots(figsize=(7,4))
    bars=ax.barh(labs,vals,color=MC_MID,height=0.5,zorder=3)
    for bar,v in zip(bars,vals):
        ax.text(bar.get_width()+0.1,bar.get_y()+bar.get_height()/2,
                str(v),va='center',fontsize=9.5,fontweight='bold',color='#333')
    ax.set_xlim(0,11.5); ax.axvline(x=5,color='#BFBFBF',linestyle='--',linewidth=0.8)
    ax.set_xlabel('Promedio (0–10)',fontsize=8,color='#595959')
    _ax_style(ax,horiz=True); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

# ── Construcción del Word ─────────────────────────────────────────────────────
# ── Construcción del Word ─────────────────────────────────────────────────────
def build_word():
    print('\n→ Generando Word...')
    doc=Document()
    for sec in doc.sections:
        sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    # ── Portada ───────────────────────────────────────────────────────────────
    tbl=doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=tbl.rows[0].cells[0]; set_cell_bg(c,C_NAVY)
    for txt,sz,bold in [
        ('INFORME DE CARACTERIZACIÓN',18,True),
        ('Monitoreo de Resultados de Tratamiento — Instrumento IRT',11,False),
        ('Ingreso a Tratamiento (IRT1)',10,False),
        (SERVICIO.upper(),14,True),(PERIODO,10,False),
        (f'N = {N} personas al ingreso',9,False),
    ]:
        p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(txt); run.font.name='Calibri'; run.font.size=Pt(sz)
        run.font.bold=bold; run.font.color.rgb=WHITE
    doc.add_paragraph()

    # ── KPIs ──────────────────────────────────────────────────────────────────
    edad_prom=R_edad.get('mean','—') if R_edad else '—'
    sust_ppal=max(R_sp,key=R_sp.get) if R_sp else '—'
    add_kpi_row(doc,[(N,'Personas ingresaron'),(f'{pct_h}%','Son hombres'),(edad_prom,'Edad promedio'),(sust_ppal,'Sust. principal')])

    # ── Presentación ──────────────────────────────────────────────────────────
    add_section_header(doc,'','Presentación')
    add_body(doc,
        f'Este informe describe el perfil de las personas que ingresaron a tratamiento '
        f'en {SERVICIO}, durante {PERIODO}, a través del instrumento IRT. '
        f'Se analizan {N} personas que completaron el IRT al ingreso (IRT1). '
        f'El {pct_h}% son hombres y el {pct_m}% son mujeres.')

    # ── Sección 1: Antecedentes ───────────────────────────────────────────────
    _page_break(doc)
    add_section_header(doc,'1','Antecedentes Generales')

    add_subsection(doc,'1.1. Distribución según Sexo')
    fig=g_sexo()
    if fig:
        buf,w=fig_to_img(fig,11); add_picture_kwnext(doc,buf,w)
    add_body(doc,f'{n_hombre} personas son hombres ({pct_h}%) y {n_mujer} son mujeres ({pct_m}%).')
    add_note(doc,f'N válido: {N_sx} personas.')

    add_subsection(doc,'1.2. Distribución según Edad')
    fig=g_edad()
    if fig:
        buf,w=fig_to_img(fig,13); add_picture_kwnext(doc,buf,w)
    if R_edad:
        gm=max(R_edad['grupos'],key=R_edad['grupos'].get)
        add_body(doc,f'Promedio de edad: {R_edad["mean"]} años (DE={R_edad["std"]}; rango {R_edad["min"]}–{R_edad["max"]}). '
                 f'Grupo más frecuente: {gm} ({R_edad["grupos"][gm]} personas).')
        add_note(doc,f'N válido: {sum(R_edad["grupos"].values())} personas.')

    # ── Sección 2: Consumo ────────────────────────────────────────────────────
    _page_break(doc)
    add_section_header(doc,'2','Consumo de Sustancias')

    add_subsection(doc,'2.1. Sustancia Principal al Ingreso')
    fig=g_sp()
    if fig:
        buf,w=fig_to_img(fig,12); add_picture_kwnext(doc,buf,w)
    if R_sp:
        ppal=max(R_sp,key=R_sp.get)
        pct_ppal=round(R_sp[ppal]/N_sp*100,1) if N_sp>0 else 0
        add_body(doc,f'La sustancia más frecuente es {ppal} ({pct_ppal}% de los casos).')
        add_note(doc,f'N válido: {N_sp} personas.')

    if dias_pp:
        add_subsection(doc,'2.2. Días de Consumo por Sustancia Principal')
        fig=g_barras(dias_pp,ylabel='Promedio días (0–28)')
        if fig:
            buf,w=fig_to_img(fig,13); add_picture_kwnext(doc,buf,w)
            add_body(doc,f'{dias_pp[0]["sust"]} tiene el mayor promedio de días de consumo: {dias_pp[0]["prom"]} días.')

    if cons_pct:
        add_subsection(doc,'2.3. % de Personas que Consumen cada Sustancia')
        fig=g_barras(cons_pct,ylabel=f'% sobre N={N}',rotacion=15)
        if fig:
            buf,w=fig_to_img(fig,13); add_picture_kwnext(doc,buf,w)
            add_body(doc,f'Los % pueden sumar más del 100% (una persona puede consumir varias). '
                     f'La más prevalente es {cons_pct[0]["sust"]} ({cons_pct[0]["pct"]}%, n={cons_pct[0]["n"]}).')

    # ── Sección 3: Salud ──────────────────────────────────────────────────────
    _page_break(doc)
    add_section_header(doc,'3','Salud')

    if n_urg is not None:
        add_subsection(doc,'3.1. Urgencias, Hospitalizaciones y Accidentes')
        add_body(doc,
            f'{n_urg} personas ({pct_urg}%) tuvieron que acudir a urgencias u hospitalizarse '
            f'como consecuencia del consumo en el mes previo al ingreso.'
            + (f' {n_acc} personas ({pct_acc}%) tuvieron algún accidente asociado.' if n_acc is not None else ''))
        add_note(doc,f'N = {N} personas.')

    if R_salud:
        add_subsection(doc,'3.2. Autopercepción del Estado de Salud')
        fig=g_salud_irt()
        if fig:
            buf,w=fig_to_img(fig,12); add_picture_kwnext(doc,buf,w)
        ps=next((d for d in R_salud if 'Psicol' in d['label']),None)
        pf=next((d for d in R_salud if 'Física' in d['label']),None)
        txt='Autopercepción del estado de salud (0=muy mala, 10=excelente).'
        if ps: txt+=f' Salud psicológica promedio: {ps["prom"]}.'
        if pf: txt+=f' Salud física promedio: {pf["prom"]}.'
        add_body(doc,txt)

    # ── Sección 4: Transgresión ───────────────────────────────────────────────
    _page_break(doc)
    add_section_header(doc,'4','Transgresión a la Norma Social')

    add_subsection(doc,'4.1. Transgresión a la Norma Social')
    fig=g_transgresion_irt()
    if fig:
        buf,w=fig_to_img(fig,11); add_picture_kwnext(doc,buf,w)
    add_body(doc,f'{n_tr} personas ({pct_tr}%) cometieron algún tipo de transgresión en el mes previo al ingreso.')
    add_note(doc,f'N = {N} personas.')

    if trans_tipos:
        add_subsection(doc,'4.2. Tipos de Transgresión')
        fig=g_barras(trans_tipos,ylabel=f'% sobre N={N}',rotacion=15)
        if fig:
            buf,w=fig_to_img(fig,13); add_picture_kwnext(doc,buf,w)
            tipo_mayor=max(trans_tipos,key=lambda x:x['pct'])
            add_body(doc,f'El tipo más frecuente es {tipo_mayor["label"]} ({tipo_mayor["pct"]}%, n={tipo_mayor["n"]}).')

    # ── Sección 5: Relaciones interpersonales ─────────────────────────────────
    if R_rel:
        _page_break(doc)
        add_section_header(doc,'5','Relaciones Interpersonales')
        add_subsection(doc,'5.1. Calidad de las Relaciones Interpersonales')
        fig=g_relaciones()
        if fig:
            buf,w=fig_to_img(fig,14); add_picture_kwnext(doc,buf,w)
        add_body(doc,'Barras azules = relaciones positivas (Excelente + Buena); grises = neutrales; rojas = negativas. '
                 'Se excluyen respuestas "No aplica".')

    # ── Sección 6: Satisfacción de vida ──────────────────────────────────────
    if R_sat:
        _page_break(doc)
        add_section_header(doc,'6','Satisfacción de Vida')
        add_subsection(doc,'6.1. Satisfacción de Vida al Ingreso')
        fig=g_satisfaccion()
        if fig:
            buf,w=fig_to_img(fig,14); add_picture_kwnext(doc,buf,w)
        mejor=max(R_sat,key=lambda x:x['prom']); peor=min(R_sat,key=lambda x:x['prom'])
        add_body(doc,f'Escala 0–10 (0=nada satisfecho, 10=muy satisfecho). '
                 f'Mayor satisfacción: {mejor["label"]} ({mejor["prom"]}). '
                 f'Menor satisfacción: {peor["label"]} ({peor["prom"]}).')

    # ── Pie ───────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p=doc.add_paragraph()
    run=p.add_run(f'Informe generado automáticamente · IRT · {SERVICIO} · {PERIODO}')
    run.font.size=Pt(8); run.font.italic=True; run.font.color.rgb=GRAY
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_FILE)
    print(f'  ✓ Word generado: {OUTPUT_FILE}')

build_word()
print(f'\n{"="*60}'); print(f'  ✅  LISTO  →  {OUTPUT_FILE}'); print(f'{"="*60}')
